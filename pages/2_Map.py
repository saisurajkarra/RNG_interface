# import streamlit as st
# import pandas as pd
# import pydeck as pdk
# from utils.api_fetch import fetch_and_merge_hh_tables
# from utils.opportunity import classify_opportunity, assign_priority
# from utils.enrich_gemini import enrich_with_gemini
# from docx import Document
# from datetime import datetime

# st.set_page_config(page_title="Map Preview & Export", layout="wide")
# st.title("🗺️ Landfill Map • Click to Preview & Save")

# # ---------------------------------------------------------------------
# # 1️⃣ Load + preprocess data (cached)
# # ---------------------------------------------------------------------
# @st.cache_data(show_spinner="Fetching & merging EPA tables...")
# def load_df():
#     df = fetch_and_merge_hh_tables()
#     df["Opportunity"] = df.apply(classify_opportunity, axis=1)
#     df["Priority"]    = df["Opportunity"].apply(assign_priority)
#     return df.dropna(subset=["latitude", "longitude"])

# df_full = load_df()
# if df_full.empty:
#     st.error("EPA API returned no facility locations.")
#     st.stop()

# # ---------------------------------------------------------------------
# # 2️⃣  User chooses subset size & optional priority filter
# # ---------------------------------------------------------------------
# with st.sidebar:
#     st.header("⚙️  Map Controls")
#     limit = st.slider("Number of facilities to plot", 50, 1000, 200, step=50)
#     prio  = st.multiselect("Show priorities", ["High", "Medium", "Low"],
#                            default=["High", "Medium", "Low"])

# df = df_full[df_full["Priority"].isin(prio)].head(limit)

# # ---------------------------------------------------------------------
# # 3️⃣  Initialise session storage for saved summaries
# # ---------------------------------------------------------------------
# if "saved_summaries" not in st.session_state:
#     st.session_state.saved_summaries = {}   # key = facility_id, value = str(summary)

# # ---------------------------------------------------------------------
# # 4️⃣  Plot map and capture click events
# # ---------------------------------------------------------------------
# INITIAL = pdk.ViewState(latitude=df["latitude"].mean(),
#                         longitude=df["longitude"].mean(),
#                         zoom=3.5)

# layer = pdk.Layer(
#     "ScatterplotLayer",
#     data=df,
#     get_position="[longitude, latitude]",
#     get_radius=6000,
#     get_fill_color="""
#         ['High','Medium','Low'].index(Priority) === 0 ? [255,0,0,160] :
#         ['High','Medium','Low'].index(Priority) === 1 ? [255,165,0,160] :
#         [0,128,0,160]
#     """,
#     pickable=True)

# r = pdk.Deck(layers=[layer], initial_view_state=INITIAL,
#              tooltip={"text": "{facility_name}\n{city}, {state}\nPriority: {Priority}"})
# event = st.pydeck_chart(r).deck_event

# # ---------------------------------------------------------------------
# # 5️⃣  If user clicked a point, enrich & save
# # ---------------------------------------------------------------------
# if event and event.get("object"):                   # user clicked marker
#     row = pd.Series(event["object"])                # convert dict to Series
#     fid = row["facility_id"]
#     st.sidebar.markdown("---")
#     st.sidebar.subheader(f"🔍 {row['facility_name']}")

#     if fid in st.session_state.saved_summaries:
#         summary = st.session_state.saved_summaries[fid]
#         st.sidebar.success("Cached")
#     else:
#         with st.spinner("Calling Gemini…"):
#             summary = enrich_with_gemini(row)
#         st.session_state.saved_summaries[fid] = summary

#     st.sidebar.text_area("Gemini Summary", summary, height=250)

# # ---------------------------------------------------------------------
# # 6️⃣  Export to DOCX
# # ---------------------------------------------------------------------
# def generate_doc(summaries: dict):
#     doc = Document()
#     doc.add_heading("Landfill Biogas Opportunity Report", 0)
#     doc.add_paragraph(f"Generated: {datetime.utcnow():%Y-%m-%d %H:%M UTC}\n")
#     for idx, (fid, text) in enumerate(summaries.items(), start=1):
#         meta = df_full[df_full["facility_id"] == fid].iloc[0]
#         doc.add_heading(f"{idx}. {meta['facility_name']} ({meta['state']})", level=1)
#         doc.add_paragraph(text)
#         doc.add_page_break()
#     path = "/mnt/data/airliquide_landfill_report.docx"
#     doc.save(path)
#     return path

# st.sidebar.markdown("---")
# if st.sidebar.button("📤 Export DOCX", disabled=not st.session_state.saved_summaries):
#     file_path = generate_doc(st.session_state.saved_summaries)
#     with open(file_path, "rb") as f:
#         st.sidebar.download_button("Download report",
#                                    data=f,
#                                    file_name="Landfill_Report.docx",
#                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
# pages/2_Map_Preview.py
# ---------------------------------------------------------------
# Interactive map → click a landfill → Gemini summary → cache
# User can enrich next 30 unseen sites per filter & export DOCX
# ---------------------------------------------------------------
# pages/2_Map_Preview.py
# ---------------------------------------------------------------
# Interactive map → click marker → Gemini summary (cached)
# Batch-enrich next 30 unseen sites per filter
# Export all cached summaries to DOCX
# ---------------------------------------------------------------

# import streamlit as st
# import pandas as pd
# import pydeck as pdk
# from datetime import datetime
# from docx import Document

# from utils.api_fetch import fetch_and_merge_hh_tables
# from utils.opportunity import classify_opportunity, assign_priority
# from utils.enrich_gemini import enrich_with_gemini, enrich_facilities_batch

# # ─────────────────────────────────────────────
# # Page config
# # ─────────────────────────────────────────────
# st.set_page_config(page_title="Landfill Map Preview", layout="wide")
# st.title("🗺️ Landfill Map • click → Gemini preview → save")

# # ─────────────────────────────────────────────
# # 1 Load + preprocess EPA data (cached)
# # ─────────────────────────────────────────────
# @st.cache_data(show_spinner="🔄 Fetching EPA landfill data …")
# def load_df():
#     df = fetch_and_merge_hh_tables()
#     df["Opportunity"] = df.apply(classify_opportunity, axis=1)
#     df["Priority"]    = df["Opportunity"].apply(assign_priority)
#     return df.dropna(subset=["latitude", "longitude"])

# df_raw = load_df()
# if df_raw.empty:
#     st.error("No location rows returned from EPA.")
#     st.stop()

# # ─────────────────────────────────────────────
# # Region helper
# # ─────────────────────────────────────────────
# REGIONS = {
#     "Northeast" : ['NY','NJ','PA','MA','CT','RI','NH','VT','ME'],
#     "Southeast" : ['FL','GA','SC','NC','AL','MS','TN','KY','VA','WV','AR'],
#     "Midwest"   : ['IL','IN','OH','MI','WI','MN','IA','MO','ND','SD','NE','KS'],
#     "West"      : ['CA','WA','OR','NV','UT','CO','AZ','NM','AK','HI','ID','MT','WY'],
# }
# def assign_region(state):
#     for r, lst in REGIONS.items():
#         if state in lst:
#             return r
#     return "Other"

# df_raw["Region"] = df_raw["state"].apply(assign_region)

# # ─────────────────────────────────────────────
# # 2 Sidebar filters + optional CSV upload
# # ─────────────────────────────────────────────
# st.sidebar.header("🔎 Filters")

# region_sel = st.sidebar.selectbox("Region", ["All"] + sorted(df_raw["Region"].unique()))
# state_sel  = st.sidebar.multiselect("States", sorted(df_raw["state"].dropna().unique()))
# prio_sel   = st.sidebar.multiselect("Priority", ["High","Medium","Low"], default=["High","Medium","Low"])
# epc_text   = st.sidebar.text_input("EPC contains")

# df = df_raw.copy()
# if region_sel != "All":
#     df = df[df["Region"] == region_sel]
# if state_sel:
#     df = df[df["state"].isin(state_sel)]
# if prio_sel:
#     df = df[df["Priority"].isin(prio_sel)]
# if epc_text:
#     df = df[df["gas_collection_sys_manufacture"].str.contains(epc_text, na=False, case=False)]

# st.success(f"{len(df):,} facilities match current filters")

# # Session cache
# if "enriched_ids" not in st.session_state:
#     st.session_state.enriched_ids = set()
#     st.session_state.enriched_records = []

# # Allow CSV import
# with st.sidebar.expander("⬆️ Upload previous enrichment CSV"):
#     up = st.file_uploader("CSV file", type="csv")
#     if up:
#         prev = pd.read_csv(up)
#         st.session_state.enriched_ids.update(prev["facility_id"])
#         st.session_state.enriched_records.extend(prev.to_dict("records"))
#         st.success(f"Merged {len(prev)} rows from file")

# # ─────────────────────────────────────────────
# # 3 Safe df for pydeck
# # ─────────────────────────────────────────────
# df_map = df[["facility_id","facility_name","state",
#              "latitude","longitude","Priority"]].copy()
# COLORS = {"High":[255,0,0,160], "Medium":[255,165,0,160], "Low":[0,128,0,160]}
# df_map["color"] = df_map["Priority"].map(COLORS)

# layer = pdk.Layer(
#     "ScatterplotLayer",
#     data=df_map,
#     get_position="[longitude, latitude]",
#     get_fill_color="color",
#     get_radius=6000,
#     pickable=True,
# )

# view = pdk.ViewState(
#     latitude=df_map["latitude"].mean(),
#     longitude=df_map["longitude"].mean(),
#     zoom=3.4,
# )

# chart = st.pydeck_chart(
#     pdk.Deck(layers=[layer], initial_view_state=view,
#              tooltip={"text":"{facility_name}\n{state}\nPriority: {Priority}"})
# )

# event = chart.deck_event       # ✅ call property to get dict

# # ─────────────────────────────────────────────
# # 4 Click marker → enrich + cache
# # ─────────────────────────────────────────────
# if isinstance(event, dict) and "object" in event:
#     row = pd.Series(event["object"])
#     fid = row["facility_id"]

#     st.sidebar.markdown("---")
#     st.sidebar.subheader(row["facility_name"])

#     if fid in st.session_state.enriched_ids:
#         summ = next(r["gemini_summary"] for r in st.session_state.enriched_records
#                     if r["facility_id"] == fid)
#         st.sidebar.success("Loaded from cache")
#     else:
#         with st.spinner("Gemini generating summary…"):
#             summ = enrich_with_gemini(row)
#         rec = row.to_dict()
#         rec["gemini_summary"] = summ
#         st.session_state.enriched_records.append(rec)
#         st.session_state.enriched_ids.add(fid)

#     st.sidebar.text_area("Gemini Summary", summ, height=280)

# # ─────────────────────────────────────────────
# # 5 Batch-enrich next 30 unseen facilities
# # ─────────────────────────────────────────────
# BATCH = 30
# queue = df[~df["facility_id"].isin(st.session_state.enriched_ids)].head(BATCH)

# st.markdown("---")
# st.markdown(f"### ⚡ Batch enrichment ({len(queue)} queued, max {BATCH})")

# if st.button("Run batch now", disabled=queue.empty):
#     texts = enrich_facilities_batch(queue, limit=len(queue))
#     for rec, txt in zip(queue.to_dict("records"), texts):
#         rec["gemini_summary"] = txt
#         st.session_state.enriched_records.append(rec)
#         st.session_state.enriched_ids.add(rec["facility_id"])
#     st.success("Batch enrichment complete!")

# # ─────────────────────────────────────────────
# # 6 Export DOCX report
# # ─────────────────────────────────────────────
# def save_docx(recs) -> str:
#     doc = Document()
#     doc.add_heading("Landfill Biogas Opportunity Report", 0)
#     doc.add_paragraph(f"Generated {datetime.utcnow():%Y-%m-%d %H:%M UTC}\n")
#     for i, r in enumerate(recs, 1):
#         doc.add_heading(f"{i}. {r['facility_name']} ({r['state']})", level=1)
#         doc.add_paragraph(r["gemini_summary"])
#         doc.add_page_break()
#     path = "/mnt/data/Landfill_Report.docx"
#     doc.save(path)
#     return path

# st.sidebar.metric("Cached summaries", len(st.session_state.enriched_records))

# if st.session_state.enriched_records:
#     if st.sidebar.button("📄 Export DOCX"):
#         path = save_docx(st.session_state.enriched_records)
#         with open(path, "rb") as f:
#             st.sidebar.download_button("Download DOCX", f,
#                 file_name="Landfill_Report.docx",
#                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# # Clear cache
# with st.sidebar.expander("🧹 Reset cache"):
#     if st.button("Clear session cache"):
#         st.session_state.enriched_ids.clear()
#         st.session_state.enriched_records.clear()
#         st.experimental_rerun()


# pages/2_Map_Preview.py
# ------------------------------------------------------------------
# Interactive map → click marker → Gemini summary (cached)
# Batch-enrich next 30 unseen; export DOCX
# Works with new Priority classes: High / Medium / Low / Very-Low / Served
# ------------------------------------------------------------------

import streamlit as st
import pandas as pd
import pydeck as pdk
from datetime import datetime
from docx import Document

from utils.api_fetch     import fetch_and_merge_hh_tables
from utils.opportunity   import classify_opportunity, assign_priority
from utils.enrich_gemini import enrich_with_gemini, enrich_facilities_batch

# ─────────────────────────────────────────────────────────────
# Page config
# ─────────────────────────────────────────────────────────────
st.set_page_config(page_title="Landfill Map Preview", layout="wide")
st.title("🗺️ Landfill Map • click → Gemini preview → save")

# ─────────────────────────────────────────────────────────────
# 1  Load & preprocess EPA data  (cached)
# ─────────────────────────────────────────────────────────────
@st.cache_data(show_spinner="🔄 Fetching EPA landfill data …")
def load_df():
    df = fetch_and_merge_hh_tables()
    df["Opportunity"] = df.apply(classify_opportunity, axis=1)
    df["Priority"]    = df["Opportunity"].apply(assign_priority)
    return df.dropna(subset=["latitude", "longitude"])

df_raw = load_df()
if df_raw.empty:
    st.error("No location rows returned from EPA.")
    st.stop()

# ─────────────────────────────────────────────────────────────
# Region helper
# ─────────────────────────────────────────────────────────────
REGIONS = {
    "Northeast": ['NY','NJ','PA','MA','CT','RI','NH','VT','ME'],
    "Southeast": ['FL','GA','SC','NC','AL','MS','TN','KY','VA','WV','AR'],
    "Midwest"  : ['IL','IN','OH','MI','WI','MN','IA','MO','ND','SD','NE','KS'],
    "West"     : ['CA','WA','OR','NV','UT','CO','AZ','NM','AK','HI','ID','MT','WY'],
}
def assign_region(state):
    for r, lst in REGIONS.items():
        if state in lst:
            return r
    return "Other"

df_raw["Region"] = df_raw["state"].apply(assign_region)

# ─────────────────────────────────────────────────────────────
# 2  Sidebar filters + optional CSV upload
# ─────────────────────────────────────────────────────────────
st.sidebar.header("🔎 Filters")

all_priorities = sorted(df_raw["Priority"].unique())
region_sel = st.sidebar.selectbox("Region", ["All"] + sorted(df_raw["Region"].unique()))
state_sel  = st.sidebar.multiselect("States", sorted(df_raw["state"].dropna().unique()))
prio_sel   = st.sidebar.multiselect("Priority", all_priorities, default=all_priorities)
epc_text   = st.sidebar.text_input("EPC contains")

df = df_raw.copy()
if region_sel != "All":
    df = df[df["Region"] == region_sel]
if state_sel:
    df = df[df["state"].isin(state_sel)]
if prio_sel:
    df = df[df["Priority"].isin(prio_sel)]
if epc_text:
    df = df[df["gas_collection_sys_manufacture"].str.contains(epc_text, na=False, case=False)]

st.success(f"{len(df):,} facilities match current filters")

# Session cache: enriched rows
if "enriched_ids" not in st.session_state:
    st.session_state.enriched_ids = set()
    st.session_state.enriched_records = []

with st.sidebar.expander("⬆️ Upload previous enrichment CSV"):
    up = st.file_uploader("CSV file", type="csv")
    if up:
        prev = pd.read_csv(up)
        st.session_state.enriched_ids.update(prev["facility_id"])
        st.session_state.enriched_records.extend(prev.to_dict("records"))
        st.success(f"Merged {len(prev)} rows from file")

# ─────────────────────────────────────────────────────────────
# 3  Safe dataframe for pydeck (avoid JS name issues)
# ─────────────────────────────────────────────────────────────
df_map = df[["facility_id","facility_name","state",
             "latitude","longitude","Priority"]].copy()

COLORS = {
    "High"     : [255,   0,   0, 160],  # red
    "Medium"   : [255, 165,   0, 160],  # orange
    "Low"      : [  0, 128,   0, 160],  # green
    "Very-Low" : [128, 128, 128, 160],  # grey
    "Served"   : [  0,   0, 255, 160],  # blue
}
default_color = [100, 100, 100, 160]
df_map["color"] = (
    df_map["Priority"]
      .map(COLORS)
      .apply(lambda c: c if isinstance(c, list) else default_color)
)

layer = pdk.Layer(
    "ScatterplotLayer",
    data=df_map,
    get_position="[longitude, latitude]",
    get_fill_color="color",
    get_radius=6000,
    pickable=True,
)

view = pdk.ViewState(
    latitude=df_map["latitude"].mean(),
    longitude=df_map["longitude"].mean(),
    zoom=3.4,
)

chart = st.pydeck_chart(
    pdk.Deck(layers=[layer], initial_view_state=view,
             tooltip={"text":"{facility_name}\n{state}\nPriority: {Priority}"}))

event = chart.deck_event   # property reference

# ─────────────────────────────────────────────────────────────
# 4  Click marker → Gemini enrich + cache
# ─────────────────────────────────────────────────────────────
if isinstance(event, dict) and "object" in event:
    row = pd.Series(event["object"])
    fid = row["facility_id"]

    st.sidebar.markdown("---")
    st.sidebar.subheader(row["facility_name"])

    if fid in st.session_state.enriched_ids:
        summ = next(r["gemini_summary"] for r in st.session_state.enriched_records
                    if r["facility_id"] == fid)
        st.sidebar.success("Loaded from cache")
    else:
        with st.spinner("Gemini generating summary…"):
            summ = enrich_with_gemini(row)
        rec = row.to_dict()
        rec["gemini_summary"] = summ
        st.session_state.enriched_records.append(rec)
        st.session_state.enriched_ids.add(fid)

    st.sidebar.text_area("Gemini Summary", summ, height=280)

# ─────────────────────────────────────────────────────────────
# 5  Batch-enrich next 30 unseen
# ─────────────────────────────────────────────────────────────
BATCH = 30
queue = df[~df["facility_id"].isin(st.session_state.enriched_ids)].head(BATCH)

st.markdown("---")
st.markdown(f"### ⚡ Batch enrichment ({len(queue)} queued, max {BATCH})")

if st.button("Run batch now", disabled=queue.empty):
    texts = enrich_facilities_batch(queue, limit=len(queue))
    for rec, txt in zip(queue.to_dict("records"), texts):
        rec["gemini_summary"] = txt
        st.session_state.enriched_records.append(rec)
        st.session_state.enriched_ids.add(rec["facility_id"])
    st.success("Batch enrichment complete!")

# ─────────────────────────────────────────────────────────────
# 6  Export DOCX
# ─────────────────────────────────────────────────────────────
def save_docx(recs):
    doc = Document()
    doc.add_heading("Landfill Biogas Opportunity Report", 0)
    doc.add_paragraph(f"Generated {datetime.utcnow():%Y-%m-%d %H:%M UTC}\n")
    for i, r in enumerate(recs, 1):
        doc.add_heading(f"{i}. {r['facility_name']} ({r['state']})", level=1)
        doc.add_paragraph(r["gemini_summary"])
        doc.add_page_break()
    path = "/mnt/data/Landfill_Report.docx"
    doc.save(path); return path

st.sidebar.metric("Cached summaries", len(st.session_state.enriched_records))

if st.session_state.enriched_records:
    if st.sidebar.button("📄 Export DOCX"):
        fp = save_docx(st.session_state.enriched_records)
        with open(fp, "rb") as f:
            st.sidebar.download_button("Download DOCX", data=f,
                file_name="Landfill_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# Clear session cache
with st.sidebar.expander("🧹 Reset cache"):
    if st.button("Clear session cache"):
        st.session_state.enriched_ids.clear()
        st.session_state.enriched_records.clear()
        st.experimental_rerun()
